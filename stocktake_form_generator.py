"""Imports"""
from datetime import datetime
import pandas
import docx
from docx.shared import Pt, Cm
#from tqdm import tqdm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Fetches current date
# Use datetime as strftime causes a bug when used with date.today()
current_date = datetime.now().strftime("%d-%b-%Y")
str_current_date = str(current_date)

csv_path = './mock-inventory.csv'
logo_path = './logo.png'

export_dataframe_csv = False

# Original columns: [0] - productcode, [2] - SumOfinstock, [6] - stockloc, [8] - description

# Reads data from csv
def data_read():
    """Reads data from CSV creates and returns dataframe"""
    # Define columns to use
    usecols = ['stockloc', 'productcode', 'description', 'SumOfinstock']

    # Read the CSV file with secific columns, skip whitespaces
    data_frame = pandas.read_csv(csv_path, encoding='utf-8', usecols=usecols, skipinitialspace = True)

    # Re-order columns - move system stock level to end
    data_frame = data_frame.reindex(columns=['stockloc', 'productcode', 'description', 'SumOfinstock'])

    # Sorts rows by stock location
    data_frame = data_frame.sort_values('stockloc')

    # Adds column with constant value. Header, Value
    data_frame['Quantity'] = ":"

    # Removes white space from description column
    data_frame['description'] = data_frame['description'].str.strip()
    return data_frame

    # Preview head and tail 5 rows
    # data.head()
    # print(data)

#======================================================================================

# Export to word docx code
def doc_builder():
    """Main document builder using dataframe as input"""
    # Initialise dataframe to use
    df = pandas.DataFrame(data_read())

    # Initialise the Word document
    doc = docx.Document()

    # Style/font settings for docx
    # Font Pt 8 creates 74 single rows per page
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(8)

    # Sets narrow margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(0.5)
        section.right_margin = Cm(0.5)

    # Create document header
    header_section = doc.sections[0]
    header = header_section.header
    header_text = header.paragraphs[0]
    header_text.text = f"\tStock-Take {str_current_date}\t"
    header_text.style = doc.styles["Heading 1"]
    header_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_text.paragraph_format.space_before = Cm(0)
    header_text.paragraph_format.space_after = Cm(0)

    # Adds logo to header
    logo_run = header_text.add_run()
    logo_run.add_picture(logo_path, width=Cm(3))

    # Initialise the table
    table = doc.add_table(rows=(df.shape[0] + 1), cols=df.shape[1])
    # Add table borders
    table.style = 'Table Grid' #'Medium List 1'
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Disable automatic sizing
    table.autofit = False
    table.allow_autofit = False

    # Set table cell size before table build loop
    # A4 page = 21cm width - 1cm for margin
    for cell in table.columns[0].cells:
        cell.width = Cm(1.5)
    # Sets the 'productcode' column width
    for cell in table.columns[1].cells:
        cell.width = Cm(5.0)
    # Sets the 'description' column width
    for cell in table.columns[2].cells:
        cell.width = Cm(9.5)
    # Sets the 'SumOfinstock' column width
    for cell in table.columns[3].cells:
        cell.width = Cm(2.0)
    # Sets the 'Quantity' column width
    for cell in table.columns[4].cells:
        cell.width = Cm(2.0)

    # Add the column headings
    for head in range(df.shape[1]):
        table.cell(0, head).text = df.columns[head]
    # Add the body of the data frame
    for i in range(df.shape[0]):
        for head, cell in enumerate(table.rows[i + 1].cells):
            cell.text = str(df.values[i, head])

    # Iterates through description column setting smaller font size
    # Progress bar within tqdm
    for cell in table.columns[2].cells[1:]:
        # Paragraphs are any value inside the cell
        paragraphs = cell.paragraphs
        for paragraph in paragraphs:
            for run in paragraph.runs:
                font = run.font
                font.size= Pt(6)

    # Export as docx
    doc.save(f'./Example-Stocktake-Form-{str_current_date}.docx')
    print("Completed.")
    # Export as new csv. No default index, Unknown for N/A values
    if export_dataframe_csv:
        df.to_csv('cleaned_database_inventory.csv', index=None, na_rep='Unknown')

#============================================================================

doc_builder()

# # Print database size
# table_size = df.shape
# print(f"Database rows (index 0) columns (index 1): {table_size}")

# # Fetches available grid style names from docx for table.style
# from docx import Document
# doc = Document()
# for style in doc.styles:
#   if "_TableStyle" in str(style):
#     print(style)
